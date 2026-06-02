from __future__ import annotations

from pathlib import Path
from typing import TypedDict

import fitz
from docx import Document
from docx.shared import Pt, RGBColor


class Section(TypedDict):
    heading: str
    level: int
    paragraphs: list[str]


CONTENT: list[Section] = [
    {
        "heading": "1. Document Parser Demo",
        "level": 1,
        "paragraphs": [
            "This is paragraph 1 under section 1 with content for extraction.",
            "This is paragraph 2 under section 1 and should be separate from paragraph 1.",
        ],
    },
    {
        "heading": "1.1 Nested Section",
        "level": 2,
        "paragraphs": [
            "This is paragraph 1 under subsection 1.1 for child heading testing.",
            "This is paragraph 2 under subsection 1.1 and should form another paragraph.",
        ],
    },
    {
        "heading": "2. Processing Notes",
        "level": 1,
        "paragraphs": [
            "Section 2 paragraph confirms multiple numbered headings are present.",
        ],
    },
]

CHALLENGING_CONTENT: list[Section] = [
    {
        "heading": "1. Executive Summary",
        "level": 1,
        "paragraphs": [
            "This opening paragraph is intentionally long to test line wrapping and paragraph reconstruction across extraction methods.",
            "This second paragraph follows after extra spacing and should remain a separate paragraph in JSON.",
        ],
    },
    {
        "heading": "1.1 Scope & Constraints",
        "level": 2,
        "paragraphs": [
            "Constraint A: Inputs may include mixed numbering schemes, punctuation, and short heading-like phrases.",
            "Constraint B: Some paragraphs contain inline markers and emphasized terms that should still map correctly.",
        ],
    },
    {
        "heading": "1.1.1 Edge Cases",
        "level": 3,
        "paragraphs": [
            "Edge case paragraph one includes dates (2026-06-02), symbols (%,$,#), and parenthetical notes (alpha, beta).",
            "Edge case paragraph two is short.",
        ],
    },
    {
        "heading": "2. Findings",
        "level": 1,
        "paragraphs": [
            "Findings paragraph one appears on a later page in PDF to test page transition behavior.",
            "Findings paragraph two concludes the challenging sample.",
        ],
    },
]

SPECIAL_PREFIX = "Special token: "
SPECIAL_MARKER = "BlueBoldItalicMarker."


def find_bold_italic_font_file() -> Path | None:
    """Return a local bold-italic font path if one is available."""
    candidates = [
        Path("/System/Library/Fonts/Supplemental/Arial Bold Italic.ttf"),
        Path("/System/Library/Fonts/Supplemental/Times New Roman Bold Italic.ttf"),
        Path("/System/Library/Fonts/Supplemental/HelveticaNeue.ttc"),
    ]
    for font_path in candidates:
        if font_path.exists() and font_path.is_file():
            return font_path
    return None


def build_docx(path: Path) -> None:
    doc = Document()
    for section in CONTENT:
        doc.add_heading(section["heading"], level=section["level"])
        for paragraph in section["paragraphs"]:
            doc.add_paragraph(paragraph)

    p = doc.add_paragraph(SPECIAL_PREFIX)
    run = p.add_run(SPECIAL_MARKER)
    run.bold = True
    run.italic = True
    run.font.color.rgb = RGBColor(0, 0, 255)

    doc.save(str(path))


def build_pdf(path: Path) -> None:
    pdf = fitz.open()
    page = pdf.new_page()

    y = 72

    for section in CONTENT:
        heading_size = 20 if section["level"] == 1 else 16
        page.insert_text((72, y), section["heading"], fontsize=heading_size, color=(0, 0, 0))
        y += 34 if section["level"] == 1 else 30
        for paragraph in section["paragraphs"]:
            page.insert_text((72, y), paragraph, fontsize=12, color=(0, 0, 0))
            y += 28
        y += 12

    # Blue text marker with explicit bold-italic font embedding when available.
    bold_italic_font = find_bold_italic_font_file()
    if bold_italic_font is not None:
        page.insert_text(
            (72, y + 28),
            SPECIAL_PREFIX + SPECIAL_MARKER,
            fontsize=12,
            color=(0, 0, 1),
            fontname="CustomBoldItalic",
            fontfile=str(bold_italic_font),
        )
    else:
        page.insert_text((72, y + 28), SPECIAL_PREFIX + SPECIAL_MARKER, fontsize=12, color=(0, 0, 1))

    pdf.save(str(path))
    pdf.close()


def build_challenging_docx(path: Path) -> None:
    doc = Document()

    for section in CHALLENGING_CONTENT:
        doc.add_heading(section["heading"], level=section["level"])
        for idx, paragraph in enumerate(section["paragraphs"], start=1):
            p = doc.add_paragraph(paragraph)
            # Add subtle per-paragraph variation to mimic real-world authoring.
            if idx == 1 and section["level"] == 1:
                p.add_run(" Additional wrapped content for extraction quality checks.")

    doc.add_heading("Appendix A. Styled Tokens", level=2)
    p = doc.add_paragraph("Token A: ")
    run_a = p.add_run(SPECIAL_MARKER)
    run_a.bold = True
    run_a.italic = True
    run_a.font.color.rgb = RGBColor(0, 0, 255)

    p2 = doc.add_paragraph("Token B: ")
    run_b = p2.add_run("BlueBoldItalicMarkerSecondary.")
    run_b.bold = True
    run_b.italic = True
    run_b.font.color.rgb = RGBColor(20, 70, 200)

    doc.add_heading("Appendix B. Rich Typography", level=2)

    p3 = doc.add_paragraph("Mixed style line: ")
    run_c1 = p3.add_run("LargeRedBold")
    run_c1.bold = True
    run_c1.font.size = Pt(16)
    run_c1.font.color.rgb = RGBColor(180, 20, 20)
    p3.add_run(" | ")
    run_c2 = p3.add_run("SmallGreenItalic")
    run_c2.italic = True
    run_c2.font.size = Pt(10)
    run_c2.font.color.rgb = RGBColor(30, 140, 60)
    p3.add_run(" | ")
    run_c3 = p3.add_run("UnderlinedBlue")
    run_c3.underline = True
    run_c3.font.color.rgb = RGBColor(0, 80, 180)

    p4 = doc.add_paragraph("Chemistry and math: H")
    sub = p4.add_run("2")
    sub.font.subscript = True
    p4.add_run("O and E = mc")
    sup = p4.add_run("2")
    sup.font.superscript = True
    p4.add_run(".")

    p5 = doc.add_paragraph("Financial notation: Total = $")
    p5_run = p5.add_run("1,250")
    p5_run.bold = True
    p5_run.font.color.rgb = RGBColor(90, 60, 10)
    p5.add_run(" (USD)")

    doc.add_heading("Appendix C. Data Table", level=2)
    table = doc.add_table(rows=4, cols=3)
    table.cell(0, 0).text = "Metric"
    table.cell(0, 1).text = "Q1"
    table.cell(0, 2).text = "Q2"
    table.cell(1, 0).text = "Revenue"
    table.cell(1, 1).text = "$1250"
    table.cell(1, 2).text = "$1430"
    table.cell(2, 0).text = "Growth"
    table.cell(2, 1).text = "6.2%"
    table.cell(2, 2).text = "7.8%"
    table.cell(3, 0).text = "Notes"
    table.cell(3, 1).text = "Pilot"
    table.cell(3, 2).text = "Scale"

    doc.save(str(path))


def _ensure_page(pdf: fitz.Document, page: fitz.Page, y: float, threshold: float = 760) -> tuple[fitz.Page, float]:
    if y <= threshold:
        return page, y
    new_page = pdf.new_page()
    return new_page, 72.0


def build_challenging_pdf(path: Path) -> None:
    pdf = fitz.open()
    page = pdf.new_page()
    y = 72.0

    for section in CHALLENGING_CONTENT:
        page, y = _ensure_page(pdf, page, y)
        heading_size = 20 if section["level"] == 1 else 16 if section["level"] == 2 else 14
        page.insert_text((72, y), section["heading"], fontsize=heading_size, color=(0, 0, 0))
        y += 34 if section["level"] == 1 else 28

        for idx, paragraph in enumerate(section["paragraphs"], start=1):
            page, y = _ensure_page(pdf, page, y)
            # Split first paragraph of each section into two nearby lines so parser should merge them.
            if idx == 1 and len(paragraph) > 85:
                first_half = paragraph[: len(paragraph) // 2].rstrip() + ""
                second_half = paragraph[len(paragraph) // 2 :].lstrip()
                page.insert_text((72, y), first_half, fontsize=12, color=(0, 0, 0))
                y += 16
                page.insert_text((72, y), second_half, fontsize=12, color=(0, 0, 0))
                y += 26
            else:
                page.insert_text((72, y), paragraph, fontsize=12, color=(0, 0, 0))
                # Add larger gap after second paragraph to force explicit paragraph boundary.
                y += 32 if idx == 2 else 24

        y += 8

    page, y = _ensure_page(pdf, page, y)
    page.insert_text((72, y), "Appendix A. Styled Tokens", fontsize=16, color=(0, 0, 0))
    y += 28

    bold_italic_font = find_bold_italic_font_file()
    styled_line_a = "Token A: " + SPECIAL_MARKER
    styled_line_b = "Token B: BlueBoldItalicMarkerSecondary."
    if bold_italic_font is not None:
        page.insert_text(
            (72, y),
            styled_line_a,
            fontsize=12,
            color=(0, 0, 1),
            fontname="CustomBoldItalicA",
            fontfile=str(bold_italic_font),
        )
        y += 22
        page.insert_text(
            (72, y),
            styled_line_b,
            fontsize=12,
            color=(0.1, 0.3, 0.85),
            fontname="CustomBoldItalicB",
            fontfile=str(bold_italic_font),
        )
    else:
        page.insert_text((72, y), styled_line_a, fontsize=12, color=(0, 0, 1))
        y += 22
        page.insert_text((72, y), styled_line_b, fontsize=12, color=(0.1, 0.3, 0.85))

    y += 34
    page, y = _ensure_page(pdf, page, y)
    page.insert_text((72, y), "Appendix B. Rich Typography", fontsize=16, color=(0, 0, 0))
    y += 28

    # Mixed font sizes/colors/styles across one visual line.
    page.insert_text((72, y), "LargeRedBold", fontsize=16, color=(0.75, 0.1, 0.1))
    page.insert_text((190, y), "SmallGreenItalic", fontsize=10, color=(0.15, 0.55, 0.2))
    page.insert_text((320, y), "UnderlinedBlue", fontsize=12, color=(0.1, 0.25, 0.8))
    y += 26

    # Subscript / superscript simulation via baseline shifts and smaller font.
    page.insert_text((72, y), "Chemistry: H", fontsize=12, color=(0, 0, 0))
    page.insert_text((145, y + 4), "2", fontsize=8, color=(0, 0, 0))
    page.insert_text((150, y), "O   Math: E = mc", fontsize=12, color=(0, 0, 0))
    page.insert_text((255, y - 4), "2", fontsize=8, color=(0, 0, 0))
    y += 24

    page.insert_text((72, y), "Financial notation: Total = $", fontsize=12, color=(0, 0, 0))
    page.insert_text((230, y), "1,250", fontsize=12, color=(0.45, 0.3, 0.1))
    page.insert_text((262, y), " (USD)", fontsize=12, color=(0, 0, 0))

    pdf.save(str(path))
    pdf.close()


def main() -> None:
    root = Path(__file__).resolve().parents[1]
    out_dir = root / "sample_inputs"
    out_dir.mkdir(parents=True, exist_ok=True)

    docx_path = out_dir / "sample_structured.docx"
    pdf_path = out_dir / "sample_structured.pdf"
    challenging_docx_path = out_dir / "sample_challenging.docx"
    challenging_pdf_path = out_dir / "sample_challenging.pdf"

    build_docx(docx_path)
    build_pdf(pdf_path)
    build_challenging_docx(challenging_docx_path)
    build_challenging_pdf(challenging_pdf_path)

    print(f"Created: {docx_path}")
    print(f"Created: {pdf_path}")
    print(f"Created: {challenging_docx_path}")
    print(f"Created: {challenging_pdf_path}")


if __name__ == "__main__":
    main()
