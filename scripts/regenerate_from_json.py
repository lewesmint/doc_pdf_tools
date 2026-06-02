from __future__ import annotations

import argparse
import json
import re
from dataclasses import dataclass
from pathlib import Path
from typing import Any, cast

import fitz
from docx import Document
from docx.shared import RGBColor


@dataclass
class Block:
    type: str
    text: str
    level: int = 1
    rows: list[list[str]] | None = None


def normalize_spaces(text: str) -> str:
    return re.sub(r"\s+", " ", text.replace("\xa0", " ")).strip()


def load_json(path: Path) -> dict[str, Any]:
    return cast(dict[str, Any], json.loads(path.read_text(encoding="utf-8")))


def flatten_blocks(node: dict[str, Any]) -> list[Block]:
    blocks: list[Block] = []

    for child in node.get("children", []):
        block_type = child.get("type", "")
        text = normalize_spaces(str(child.get("text", "")))
        if not text:
            continue

        if block_type == "heading":
            level = int(child.get("level", 1) or 1)
            blocks.append(Block(type="heading", text=text, level=max(1, min(level, 9))))
        elif block_type == "paragraph":
            blocks.append(Block(type="paragraph", text=text, level=1))
        elif block_type == "table":
            rows_raw = child.get("rows", [])
            rows: list[list[str]] = []
            if isinstance(rows_raw, list):
                for row in rows_raw:
                    if not isinstance(row, list):
                        continue
                    rows.append([normalize_spaces(str(cell)) for cell in row])
            if rows:
                blocks.append(Block(type="table", text="table", level=1, rows=rows))

        blocks.extend(flatten_blocks(child))

    return blocks


def special_markers(data: dict[str, Any]) -> list[str]:
    markers: list[str] = []
    for item in data.get("special_formatted_text", []):
        text = normalize_spaces(str(item.get("text", "")))
        if text:
            markers.append(text)
    # Longest first prevents shorter markers from breaking longer matches.
    return sorted(set(markers), key=len, reverse=True)


def special_style_map(data: dict[str, Any]) -> dict[str, dict[str, Any]]:
    styles: dict[str, dict[str, Any]] = {}
    for item in data.get("special_formatted_text", []):
        text = normalize_spaces(str(item.get("text", "")))
        if not text:
            continue
        style = item.get("style", {}) if isinstance(item, dict) else {}
        styles[text] = style if isinstance(style, dict) else {}
    return styles


def style_for_segment(
    part_text: str,
    style_map: dict[str, dict[str, Any]],
) -> dict[str, Any] | None:
    key = normalize_spaces(part_text)
    return style_map.get(key)


def style_for_paragraph(text: str, style_map: dict[str, dict[str, Any]]) -> dict[str, Any] | None:
    text_norm = normalize_spaces(text)
    for marker, style in style_map.items():
        if marker in text_norm:
            return style
    return None


def find_bold_italic_font_file() -> Path | None:
    candidates = [
        Path("/System/Library/Fonts/Supplemental/Arial Bold Italic.ttf"),
        Path("/System/Library/Fonts/Supplemental/Times New Roman Bold Italic.ttf"),
        Path("/System/Library/Fonts/Supplemental/HelveticaNeue.ttc"),
    ]
    for font_path in candidates:
        if font_path.exists() and font_path.is_file():
            return font_path
    return None


def split_by_markers(text: str, markers: list[str]) -> list[tuple[str, bool]]:
    parts: list[tuple[str, bool]] = [(text, False)]
    for marker in markers:
        next_parts: list[tuple[str, bool]] = []
        marker_norm = normalize_spaces(marker)
        for part_text, is_special in parts:
            if is_special:
                next_parts.append((part_text, is_special))
                continue

            idx = normalize_spaces(part_text).find(marker_norm)
            if idx < 0:
                next_parts.append((part_text, False))
                continue

            # Use direct string split for exact placement in the original paragraph text.
            raw_idx = part_text.find(marker)
            if raw_idx < 0:
                raw_idx = part_text.find(marker_norm)
            if raw_idx < 0:
                next_parts.append((part_text, False))
                continue

            before = part_text[:raw_idx]
            match = part_text[raw_idx : raw_idx + len(marker)]
            after = part_text[raw_idx + len(marker) :]
            if before:
                next_parts.append((before, False))
            if match:
                next_parts.append((match, True))
            if after:
                next_parts.append((after, False))

        parts = next_parts
    return parts


def regenerate_docx(data: dict[str, Any], output_path: Path) -> None:
    doc = Document()
    blocks = flatten_blocks(data.get("structure", {}))
    markers = special_markers(data)
    styles = special_style_map(data)

    for block in blocks:
        if block.type == "heading":
            doc.add_heading(block.text, level=max(1, min(block.level, 9)))
            continue

        if block.type == "table":
            table_rows = block.rows or []
            if not table_rows:
                continue
            col_count = max(len(r) for r in table_rows)
            table = doc.add_table(rows=len(table_rows), cols=col_count)
            for r_idx, row in enumerate(table_rows):
                for c_idx in range(col_count):
                    cell_text = row[c_idx] if c_idx < len(row) else ""
                    table.cell(r_idx, c_idx).text = cell_text
            continue

        paragraph = doc.add_paragraph()
        for part_text, is_special in split_by_markers(block.text, markers):
            run = paragraph.add_run(part_text)
            if is_special:
                style = style_for_segment(part_text, styles) or {}
                run.bold = bool(style.get("bold", True))
                run.italic = bool(style.get("italic", True))
                run.underline = bool(style.get("underline", False))
                run.font.subscript = bool(style.get("subscript", False))
                run.font.superscript = bool(style.get("superscript", False))

                rgb = (
                    style.get("color_rgb", {})
                    if isinstance(style.get("color_rgb", {}), dict)
                    else {}
                )
                r = int(rgb.get("r", 0))
                g = int(rgb.get("g", 0))
                b = int(rgb.get("b", 255))
                run.font.color.rgb = RGBColor(r, g, b)

    doc.save(str(output_path))


def estimate_box_height(text: str, width: float, fontsize: float) -> float:
    chars_per_line = max(20, int(width / (fontsize * 0.5)))
    lines = max(1, (len(text) // chars_per_line) + 1)
    return lines * (fontsize * 1.35)


def regenerate_pdf(data: dict[str, Any], output_path: Path) -> None:
    doc = fitz.open()
    page = doc.new_page()

    margin_left = 72.0
    margin_top = 72.0
    page_bottom = 760.0
    usable_width = 468.0
    y = margin_top

    blocks = flatten_blocks(data.get("structure", {}))
    styles = special_style_map(data)
    bi_font_file = find_bold_italic_font_file()

    def ensure_page(required_height: float) -> None:
        nonlocal page, y
        if y + required_height <= page_bottom:
            return
        page = doc.new_page()
        y = margin_top

    for block in blocks:
        if block.type == "heading":
            fontsize = {1: 20.0, 2: 16.0, 3: 14.0}.get(block.level, 13.0)
            line_h = (fontsize * 1.45) + 6
            ensure_page(line_h)
            # Direct heading rendering is more stable than textbox for preserving heading lines.
            page.insert_text((margin_left, y), block.text, fontsize=fontsize, color=(0, 0, 0))
            y += line_h + 6
            continue

        style = style_for_paragraph(block.text, styles)
        is_special = style is not None
        style = style or {}
        fontsize = float(style.get("font_size", 12.0) or 12.0)
        box_h = estimate_box_height(block.text, usable_width, fontsize) + 6
        ensure_page(box_h)
        rect = fitz.Rect(margin_left, y, margin_left + usable_width, y + box_h)

        rgb = style.get("color_rgb", {}) if isinstance(style.get("color_rgb", {}), dict) else {}
        r = int(rgb.get("r", 0))
        g = int(rgb.get("g", 0))
        b = int(rgb.get("b", 0))
        color = (r / 255.0, g / 255.0, b / 255.0)

        is_bold = bool(style.get("bold", False))
        is_italic = bool(style.get("italic", False))

        if is_special and is_bold and is_italic and bi_font_file is not None:
            page.insert_textbox(
                rect,
                block.text,
                fontsize=fontsize,
                color=color,
                fontname="RoundTripBoldItalic",
                fontfile=str(bi_font_file),
                align=fitz.TEXT_ALIGN_LEFT,
            )
        elif is_special:
            page.insert_textbox(
                rect,
                block.text,
                fontsize=fontsize,
                color=color,
                align=fitz.TEXT_ALIGN_LEFT,
            )
        else:
            page.insert_textbox(
                rect,
                block.text,
                fontsize=fontsize,
                color=(0, 0, 0),
                align=fitz.TEXT_ALIGN_LEFT,
            )

        y += box_h + 10

    doc.save(str(output_path))
    doc.close()


def main() -> None:
    parser = argparse.ArgumentParser(description="Regenerate DOCX/PDF files from extracted JSON.")
    parser.add_argument("input_json", type=Path, help="Path to extracted JSON")
    parser.add_argument("--output", type=Path, default=None, help="Output file path")
    parser.add_argument(
        "--format",
        choices=["docx", "pdf"],
        default=None,
        help="Output format. Defaults to the source type in JSON.",
    )
    args = parser.parse_args()

    data = load_json(args.input_json)
    src_type = str(data.get("type", "")).lower()
    out_format = args.format or src_type
    if out_format not in {"docx", "pdf"}:
        raise ValueError("Could not infer output format. Use --format docx|pdf.")

    if args.output is not None:
        output_path = args.output
    else:
        output_path = args.input_json.with_suffix(f".roundtrip.{out_format}")

    output_path.parent.mkdir(parents=True, exist_ok=True)

    if out_format == "docx":
        regenerate_docx(data, output_path)
    else:
        regenerate_pdf(data, output_path)

    print(f"Wrote regenerated file: {output_path}")


if __name__ == "__main__":
    main()
