from __future__ import annotations

from pathlib import Path

import fitz
import pytest
from docx import Document

from extract_structure import extract_file, parse_docx, parse_pdf


def test_parse_docx_empty_document_returns_empty_structure(tmp_path: Path) -> None:
    docx_path = tmp_path / "empty.docx"
    Document().save(str(docx_path))

    result = parse_docx(docx_path)

    assert result["type"] == "docx"
    assert result["structure"]["type"] == "document"
    assert result["structure"].get("children", []) == []
    assert not result["special_formatted_text"]


def test_parse_pdf_empty_page_returns_empty_structure(tmp_path: Path) -> None:
    pdf_path = tmp_path / "empty.pdf"
    pdf = fitz.open()
    pdf.new_page()
    pdf.save(str(pdf_path))
    pdf.close()

    result = parse_pdf(pdf_path)

    assert result["type"] == "pdf"
    assert result["structure"]["type"] == "document"
    assert result["structure"].get("children", []) == []
    assert not result["special_formatted_text"]


def test_parse_docx_without_headings_keeps_paragraphs_at_root(tmp_path: Path) -> None:
    docx_path = tmp_path / "no_headings.docx"
    doc = Document()
    doc.add_paragraph("Paragraph one")
    doc.add_paragraph("Paragraph two")
    doc.save(str(docx_path))

    result = parse_docx(docx_path)

    children = result["structure"]["children"]
    assert [node["type"] for node in children] == ["paragraph", "paragraph"]
    assert [node["text"] for node in children] == ["Paragraph one", "Paragraph two"]


def test_extract_file_rejects_legacy_doc(tmp_path: Path) -> None:
    doc_path = tmp_path / "legacy.doc"
    doc_path.write_text("placeholder", encoding="utf-8")

    with pytest.raises(ValueError, match="Legacy \\.doc files"):
        extract_file(doc_path)


def test_extract_file_rejects_unsupported_extension(tmp_path: Path) -> None:
    txt_path = tmp_path / "notes.txt"
    txt_path.write_text("placeholder", encoding="utf-8")

    with pytest.raises(ValueError, match="Unsupported file type"):
        extract_file(txt_path)
