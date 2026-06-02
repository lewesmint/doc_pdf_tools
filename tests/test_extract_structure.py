from __future__ import annotations

from typing import Any

import fitz
from docx import Document
from docx.shared import RGBColor

from extract_structure import is_blue_rgb, parse_docx, parse_pdf


def _collect_nodes_by_type(node: dict[str, Any], node_type: str) -> list[dict[str, Any]]:
    results: list[dict[str, Any]] = []
    if node.get("type") == node_type:
        results.append(node)
    for child in node.get("children", []):
        results.extend(_collect_nodes_by_type(child, node_type))
    return results


def test_is_blue_rgb_thresholds() -> None:
    assert is_blue_rgb(0, 0, 255)
    assert is_blue_rgb(20, 50, 180)
    assert not is_blue_rgb(120, 120, 120)
    assert not is_blue_rgb(200, 20, 30)


def test_parse_docx_extracts_hierarchy_and_special_text(tmp_path) -> None:
    doc = Document()
    doc.add_heading("Main Heading", level=1)
    doc.add_paragraph("Body paragraph under heading.")
    doc.add_heading("Child Heading", level=2)

    formatted_para = doc.add_paragraph()
    normal_run = formatted_para.add_run("RegularRun ")
    normal_run.italic = True
    special_run = formatted_para.add_run("BlueBoldItalicDocx")
    special_run.bold = True
    special_run.italic = True
    special_run.font.color.rgb = RGBColor(0, 0, 255)

    docx_path = tmp_path / "sample.docx"
    doc.save(str(docx_path))

    result = parse_docx(docx_path)

    assert result["type"] == "docx"
    structure = result["structure"]
    assert structure["type"] == "document"

    headings = _collect_nodes_by_type(structure, "heading")
    heading_texts = [h["text"] for h in headings]
    assert "Main Heading" in heading_texts
    assert "Child Heading" in heading_texts

    paragraphs = _collect_nodes_by_type(structure, "paragraph")
    paragraph_texts = [p["text"] for p in paragraphs]
    assert any("Body paragraph under heading." in text for text in paragraph_texts)

    extracted_runs = result["extracted_text_runs"]
    assert any(item["text"] == "RegularRun " for item in extracted_runs)
    assert any(item["text"] == "BlueBoldItalicDocx" for item in extracted_runs)
    regular_run = next(item for item in extracted_runs if item["text"] == "RegularRun ")
    special_run = next(item for item in extracted_runs if item["text"] == "BlueBoldItalicDocx")
    regular_style = regular_run["style"]
    special_style = special_run["style"]
    assert regular_style["italic"] is True
    assert regular_style["bold"] is False
    assert special_style["bold"] is True
    assert special_style["italic"] is True
    assert special_style["color_rgb"] == {"r": 0, "g": 0, "b": 255}
    assert regular_run["char_start"] == 0
    assert regular_run["char_end"] == len("RegularRun ")
    assert special_run["char_start"] == len("RegularRun ")
    assert special_run["char_end"] == len("RegularRun BlueBoldItalicDocx")

    specials = result["special_formatted_text"]
    assert any(item["text"] == "BlueBoldItalicDocx" for item in specials)


def test_parse_pdf_extracts_heading_and_paragraph_text(tmp_path) -> None:
    pdf_path = tmp_path / "sample.pdf"

    pdf = fitz.open()
    page = pdf.new_page()
    page.insert_text((72, 72), "Main Heading", fontsize=20, color=(0, 0, 0))
    page.insert_text(
        (72, 110),
        "This is body paragraph content.",
        fontsize=12,
        fontname="helv",
        color=(0, 0, 0),
    )
    page.insert_text(
        (72, 130),
        "Another body line to anchor body font inference.",
        fontsize=12,
        fontname="helv",
        color=(0, 0, 0),
    )
    pdf.save(str(pdf_path))
    pdf.close()

    result = parse_pdf(pdf_path)

    assert result["type"] == "pdf"
    structure = result["structure"]

    headings = _collect_nodes_by_type(structure, "heading")
    heading_texts = [h["text"] for h in headings]
    assert any("Main Heading" in text for text in heading_texts)

    paragraphs = _collect_nodes_by_type(structure, "paragraph")
    paragraph_texts = [p["text"] for p in paragraphs]
    assert any("This is body paragraph content." in text for text in paragraph_texts)

    extracted_runs = result["extracted_text_runs"]
    assert any(item["text"] == "Main Heading" for item in extracted_runs)
    assert any(item["text"] == "This is body paragraph content." for item in extracted_runs)

    assert isinstance(result["special_formatted_text"], list)


def test_parse_pdf_preserves_span_text_and_offsets(tmp_path) -> None:
    pdf_path = tmp_path / "span_offsets.pdf"

    pdf = fitz.open()
    page = pdf.new_page()
    page.insert_text((72, 100), "Alpha", fontsize=12, fontname="helv", color=(0, 0, 0))
    page.insert_text((105, 100), " ", fontsize=12, fontname="helv", color=(0, 0, 0))
    page.insert_text((112, 100), "Beta", fontsize=12, fontname="helv", color=(1, 0, 0))
    pdf.save(str(pdf_path))
    pdf.close()

    result = parse_pdf(pdf_path)
    runs = [r for r in result["extracted_text_runs"] if r["page"] == 1]
    runs = sorted(runs, key=lambda r: r["char_start"])

    assert len(runs) >= 2
    assert all(r["line_index"] == runs[0]["line_index"] for r in runs)
    assert runs[0]["char_start"] == 0
    assert all(runs[idx]["char_end"] == runs[idx + 1]["char_start"] for idx in range(len(runs) - 1))

    reconstructed = "".join(r["text"] for r in runs)
    assert reconstructed.startswith("Alpha")
    assert reconstructed.endswith("Beta")


def test_parse_docx_extracts_tables(tmp_path) -> None:
    docx_path = tmp_path / "table_sample.docx"
    doc = Document()
    doc.add_heading("Data", level=1)
    table = doc.add_table(rows=3, cols=2)
    table.cell(0, 0).text = "Metric"
    table.cell(0, 1).text = "Value"
    table.cell(1, 0).text = "Revenue"
    table.cell(1, 1).text = "$1250"
    table.cell(2, 0).text = "Growth"
    table.cell(2, 1).text = "6.2%"
    doc.save(str(docx_path))

    result = parse_docx(docx_path)
    tables = _collect_nodes_by_type(result["structure"], "table")

    assert len(tables) == 1
    rows = tables[0]["rows"]
    assert rows[0] == ["Metric", "Value"]
    assert rows[1] == ["Revenue", "$1250"]
    assert rows[2] == ["Growth", "6.2%"]
