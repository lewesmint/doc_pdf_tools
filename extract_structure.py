#!/usr/bin/env python3
"""Extract structured content from PDF or DOCX into JSON.

Features:
- Heading hierarchy (including child headings)
- Paragraph content grouped under nearest heading
- Specially formatted text detection (blue + italic + bold)
"""

from __future__ import annotations

import argparse
import json
from collections import Counter
from collections.abc import Iterable
from dataclasses import dataclass, field
from pathlib import Path
from typing import Any, cast

import fitz  # PyMuPDF
from docx import Document
from docx.document import Document as DocxDocument
from docx.oxml.table import CT_Tbl
from docx.oxml.text.paragraph import CT_P
from docx.table import Table, _Cell
from docx.text.paragraph import Paragraph


@dataclass
class Node:
    """Tree node representing document structure content."""

    type: str
    text: str
    level: int | None = None
    rows: list[list[str]] | None = None
    children: list[Node] = field(default_factory=list)

    def to_dict(self) -> dict[str, Any]:
        data: dict[str, Any] = {"type": self.type, "text": self.text}
        if self.level is not None:
            data["level"] = self.level
        if self.rows is not None:
            data["rows"] = self.rows
        if self.children:
            data["children"] = [c.to_dict() for c in self.children]
        return data


def is_blue_rgb(r: int, g: int, b: int) -> bool:
    # Allow near-blue values because source files may use theme shades.
    return r <= 80 and g <= 100 and b >= 120 and b >= r + 30 and b >= g + 20


def heading_stack_push(root: Node, stack: list[Node], heading_text: str, level: int) -> Node:
    heading = Node(type="heading", text=heading_text, level=level)
    while stack and (stack[-1].level or 0) >= level:
        stack.pop()
    parent = stack[-1] if stack else root
    parent.children.append(heading)
    stack.append(heading)
    return heading


def add_paragraph(root: Node, stack: list[Node], text: str) -> None:
    text = text.strip()
    if not text:
        return
    parent = stack[-1] if stack else root
    parent.children.append(Node(type="paragraph", text=text))


def add_table(root: Node, stack: list[Node], rows: list[list[str]]) -> None:
    if not rows:
        return
    parent = stack[-1] if stack else root
    parent.children.append(Node(type="table", text="table", rows=rows))


def normalize_docx_cell_text(cell: _Cell) -> str:
    parts = [p.text.strip() for p in cell.paragraphs if p.text and p.text.strip()]
    return " ".join(parts).strip()


def iter_block_items(parent: DocxDocument | _Cell) -> Iterable[Paragraph | Table]:
    """Yield paragraphs and tables in document order."""
    parent_elm = parent.element.body if isinstance(parent, DocxDocument) else parent._tc  # pylint: disable=protected-access
    for child in parent_elm.iterchildren():
        if isinstance(child, CT_P):
            yield Paragraph(child, parent)
        elif isinstance(child, CT_Tbl):
            yield Table(child, parent)


def collect_special_docx_runs(
    para: Paragraph,
    paragraph_index: int,
    special_text_runs: list[dict[str, Any]],
) -> None:
    for r_idx, run in enumerate(para.runs):
        run_text = run.text.strip()
        if not run_text:
            continue

        is_bold = bool(run.bold)
        is_italic = bool(run.italic)
        is_underline = bool(run.underline)
        is_subscript = bool(run.font.subscript)
        is_superscript = bool(run.font.superscript)

        rgb = None
        if run.font.color is not None and run.font.color.rgb is not None:
            rgb_hex = str(run.font.color.rgb)
            if len(rgb_hex) == 6:
                rgb = (
                    int(rgb_hex[0:2], 16),
                    int(rgb_hex[2:4], 16),
                    int(rgb_hex[4:6], 16),
                )

        if rgb and is_bold and is_italic and is_blue_rgb(*rgb):
            special_text_runs.append(
                {
                    "text": run_text,
                    "page": None,
                    "paragraph_index": paragraph_index,
                    "run_index": r_idx,
                    "style": {
                        "bold": is_bold,
                        "italic": is_italic,
                        "underline": is_underline,
                        "subscript": is_subscript,
                        "superscript": is_superscript,
                        "color_rgb": {"r": rgb[0], "g": rgb[1], "b": rgb[2]},
                    },
                }
            )


def parse_docx(path: Path) -> dict[str, Any]:
    doc = Document(str(path))
    root = Node(type="document", text=path.name)
    heading_stack: list[Node] = []
    special_text_runs: list[dict[str, Any]] = []
    paragraph_index = 0

    for block in iter_block_items(doc):
        if isinstance(block, Paragraph):
            p_text = block.text.strip()

            style_name = ""
            if block.style is not None and block.style.name is not None:
                style_name = block.style.name.strip().lower()
            level = None
            if style_name.startswith("heading"):
                suffix = style_name.replace("heading", "").strip()
                if suffix.isdigit():
                    level = int(suffix)
                else:
                    level = 1

            if p_text:
                if level is not None:
                    heading_stack_push(root, heading_stack, p_text, level)
                else:
                    add_paragraph(root, heading_stack, p_text)

            collect_special_docx_runs(block, paragraph_index, special_text_runs)
            paragraph_index += 1
            continue

        # Table block
        rows: list[list[str]] = []
        for row in block.rows:
            row_cells: list[str] = []
            for cell in row.cells:
                row_cells.append(normalize_docx_cell_text(cell))
                for cell_paragraph in cell.paragraphs:
                    collect_special_docx_runs(cell_paragraph, paragraph_index, special_text_runs)
                    paragraph_index += 1
            rows.append(row_cells)
        add_table(root, heading_stack, rows)

    return {
        "source_file": str(path),
        "type": "docx",
        "structure": root.to_dict(),
        "special_formatted_text": special_text_runs,
    }


def span_flags_to_styles(flags: int, font_name: str) -> tuple[bool, bool]:
    font_lower = font_name.lower()
    is_italic = bool(flags & 2) or "italic" in font_lower or "oblique" in font_lower
    is_bold = bool(flags & 16) or "bold" in font_lower
    return is_bold, is_italic


def infer_body_font_size(lines: list[dict[str, Any]]) -> float:
    sizes = [round(line["max_size"], 1) for line in lines if line["text"].strip()]
    if not sizes:
        return 12.0
    # Mode is robust for body text in most documents.
    count = Counter(sizes)
    return float(count.most_common(1)[0][0])


def infer_heading_levels(lines: list[dict[str, Any]], body_size: float) -> dict[float, int]:
    heading_sizes: list[float] = []
    for line in lines:
        text = line["text"].strip()
        if not text:
            continue
        max_size = round(line["max_size"], 1)
        is_bold_line = line["has_bold"]
        likely_heading = max_size >= body_size + 1.0 or (
            is_bold_line and len(text) <= 100 and not text.endswith((".", ":", ";", ","))
        )
        if likely_heading:
            heading_sizes.append(max_size)

    if not heading_sizes:
        return {}

    unique_sizes = sorted(set(heading_sizes), reverse=True)
    return {size: idx + 1 for idx, size in enumerate(unique_sizes[:6])}


def parse_pdf(path: Path) -> dict[str, Any]:
    doc = fitz.open(path)
    root = Node(type="document", text=path.name)
    heading_stack: list[Node] = []
    special_text_runs: list[dict[str, Any]] = []

    extracted_lines: list[dict[str, Any]] = []

    for page_idx in range(len(doc)):
        page_number = page_idx + 1
        page = doc[page_idx]
        page_dict = cast(dict[str, Any], page.get_text("dict"))
        blocks = cast(list[dict[str, Any]], page_dict.get("blocks", []))

        for block in blocks:
            if block.get("type") != 0:
                continue
            lines = cast(list[dict[str, Any]], block.get("lines", []))
            for line in lines:
                spans = cast(list[dict[str, Any]], line.get("spans", []))
                texts: list[str] = []
                max_size = 0.0
                has_bold = False
                y_top = None

                for span in spans:
                    text = span.get("text", "")
                    if not text:
                        continue
                    text = text.replace("\xa0", " ")
                    texts.append(text)

                    size = float(span.get("size", 0.0))
                    max_size = max(max_size, size)

                    flags = int(span.get("flags", 0))
                    font = span.get("font", "")
                    is_bold, is_italic = span_flags_to_styles(flags, font)
                    has_bold = has_bold or is_bold

                    color_int = int(span.get("color", 0))
                    r = (color_int >> 16) & 255
                    g = (color_int >> 8) & 255
                    b = color_int & 255
                    if is_bold and is_italic and is_blue_rgb(r, g, b):
                        special_text_runs.append(
                            {
                                "text": text.strip().replace("\xa0", " "),
                                "page": page_number,
                                "paragraph_index": None,
                                "run_index": None,
                                "style": {
                                    "bold": is_bold,
                                    "italic": is_italic,
                                    "color_rgb": {"r": r, "g": g, "b": b},
                                    "font": font,
                                    "font_size": size,
                                },
                            }
                        )

                    bbox = span.get("bbox")
                    if bbox and y_top is None:
                        y_top = float(bbox[1])

                full_text = "".join(texts).strip()
                if not full_text:
                    continue

                extracted_lines.append(
                    {
                        "page": page_number,
                        "text": full_text,
                        "max_size": max_size,
                        "has_bold": has_bold,
                        "y_top": y_top if y_top is not None else 0.0,
                    }
                )

    body_size = infer_body_font_size(extracted_lines)
    heading_level_by_size = infer_heading_levels(extracted_lines, body_size)

    paragraph_buffer: list[str] = []
    last_body_page: int | None = None
    last_body_y_top: float | None = None
    paragraph_break_gap = body_size * 1.8

    def flush_paragraph_buffer() -> None:
        if paragraph_buffer:
            add_paragraph(root, heading_stack, " ".join(paragraph_buffer).strip())
            paragraph_buffer.clear()

    for line in extracted_lines:
        text = line["text"].strip()
        max_size = round(line["max_size"], 1)
        is_bold_line = line["has_bold"]

        is_heading = False
        level = None
        if max_size in heading_level_by_size:
            if max_size >= body_size + 1.0 or (
                is_bold_line and len(text) <= 100 and not text.endswith((".", ":", ";", ","))
            ):
                is_heading = True
                level = heading_level_by_size[max_size]

        if is_heading and level is not None:
            flush_paragraph_buffer()
            heading_stack_push(root, heading_stack, text, level)
            last_body_page = None
            last_body_y_top = None
        else:
            current_page = int(line["page"])
            current_y_top = float(line["y_top"])
            should_break_paragraph = False

            if paragraph_buffer and last_body_page is not None and last_body_y_top is not None:
                page_changed = current_page != last_body_page
                vertical_gap = current_y_top - last_body_y_top
                if page_changed or vertical_gap >= paragraph_break_gap:
                    should_break_paragraph = True

            if should_break_paragraph:
                flush_paragraph_buffer()

            paragraph_buffer.append(text)
            last_body_page = current_page
            last_body_y_top = current_y_top

    flush_paragraph_buffer()

    return {
        "source_file": str(path),
        "type": "pdf",
        "structure": root.to_dict(),
        "special_formatted_text": [s for s in special_text_runs if s["text"]],
        "pdf_inference": {
            "body_font_size": body_size,
            "heading_levels_by_font_size": heading_level_by_size,
            "note": (
                "PDF heading detection is heuristic and may need threshold "
                "tuning for specific templates."
            ),
        },
    }


def extract_file(input_path: Path) -> dict[str, Any]:
    suffix = input_path.suffix.lower()
    if suffix == ".docx":
        return parse_docx(input_path)
    if suffix == ".pdf":
        return parse_pdf(input_path)
    if suffix == ".doc":
        raise ValueError(
            "Legacy .doc files are not directly supported in this script. "
            "Convert .doc to .docx first."
        )
    raise ValueError(f"Unsupported file type: {suffix}. Use .pdf or .docx")


def main() -> None:
    parser = argparse.ArgumentParser(
        description=(
            "Extract heading hierarchy, paragraphs, and special formatting "
            "from PDF/DOCX into JSON."
        )
    )
    parser.add_argument("input", type=Path, help="Path to source PDF or DOCX")
    parser.add_argument(
        "-o",
        "--output",
        type=Path,
        default=None,
        help="Output JSON file path. Defaults to <input>.json",
    )
    args = parser.parse_args()

    input_path: Path = args.input
    if not input_path.exists():
        raise FileNotFoundError(f"Input file not found: {input_path}")

    result = extract_file(input_path)

    output_path = args.output or input_path.with_suffix(input_path.suffix + ".json")
    output_path.write_text(json.dumps(result, indent=2, ensure_ascii=False), encoding="utf-8")
    print(f"Wrote structured JSON to: {output_path}")


if __name__ == "__main__":
    main()
